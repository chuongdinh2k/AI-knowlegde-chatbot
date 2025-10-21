from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import List, Optional
import uuid
import os
from app.database import get_db, Document
from app.models import DocumentResponse, DocumentListResponse, ErrorResponse
from app.services.document_service import document_service
import logging

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["documents"])

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Upload and process a document"""
    try:
        # Validate file type
        allowed_types = ['.txt', '.pdf', '.docx', '.md']
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension not in allowed_types:
            raise HTTPException(
                status_code=400,
                detail=f"File type {file_extension} not supported. Allowed types: {allowed_types}"
            )
        
        # Read file content
        content = await file.read()
        
        # Convert to text based on file type
        if file_extension == '.txt' or file_extension == '.md':
            text_content = content.decode('utf-8')
        elif file_extension == '.pdf':
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text()
        elif file_extension == '.docx':
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(content))
            text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            text_content = content.decode('utf-8', errors='ignore')
        
        # Create document record
        document = Document(
            filename=file.filename,
            content=text_content,
            file_type=file_extension,
            file_size=len(content),
            metadata=f'{{"original_filename": "{file.filename}", "upload_method": "api"}}'
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # Process document asynchronously (in production, use a task queue)
        try:
            document_service.process_document(db, str(document.id), text_content)
        except Exception as e:
            logger.error(f"Error processing document {document.id}: {e}")
            # Document is saved but not processed
        
        return DocumentResponse(
            id=document.id,
            filename=document.filename,
            file_type=document.file_type,
            file_size=document.file_size,
            upload_date=document.upload_date,
            processed=document.processed,
            metadata=document.metadata
        )
        
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/", response_model=DocumentListResponse)
async def list_documents(
    skip: int = 0,
    limit: int = 20,
    db: Session = Depends(get_db)
):
    """List all uploaded documents"""
    try:
        documents = db.query(Document).offset(skip).limit(limit).all()
        total = db.query(Document).count()
        
        return DocumentListResponse(
            documents=[
                DocumentResponse(
                    id=doc.id,
                    filename=doc.filename,
                    file_type=doc.file_type,
                    file_size=doc.file_size,
                    upload_date=doc.upload_date,
                    processed=doc.processed,
                    metadata=doc.metadata
                )
                for doc in documents
            ],
            total=total
        )
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: str,
    db: Session = Depends(get_db)
):
    """Get a specific document by ID"""
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        return DocumentResponse(
            id=document.id,
            filename=document.filename,
            file_type=document.file_type,
            file_size=document.file_size,
            upload_date=document.upload_date,
            processed=document.processed,
            metadata=document.metadata
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/{document_id}")
async def delete_document(
    document_id: str,
    db: Session = Depends(get_db)
):
    """Delete a document and all its chunks"""
    try:
        success = document_service.delete_document(db, document_id)
        if not success:
            raise HTTPException(status_code=404, detail="Document not found or could not be deleted")
        
        return {"message": "Document deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/{document_id}/reprocess")
async def reprocess_document(
    document_id: str,
    db: Session = Depends(get_db)
):
    """Reprocess a document to update its chunks and embeddings"""
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        success = document_service.process_document(db, document_id, document.content)
        if not success:
            raise HTTPException(status_code=500, detail="Failed to reprocess document")
        
        return {"message": "Document reprocessed successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error reprocessing document: {e}")
        raise HTTPException(status_code=500, detail=str(e))
